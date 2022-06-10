# -*- coding: utf-8 -*-
"""
*******************************************************************************
 * @fileName : get_text_from_docx.py
 * @author   : "Ko Sun Ho"
 * @comment  : 문서 파싱 docx module(lib - python-docx)
 * @revision history
 * date            author         comment
 * ----------      ---------      ----------------------------------------------
 * 2022. 05. 16.  Ko Sun Ho       작성
 *******************************************************************************
"""
import docx


class GetDocxText:
    def __init__(self, filepath, filename):
        self.filepath = filepath
        self.filename = filename
        self.fullpath = self.filepath + '/' + self.filename + '.docx'
        self.doc = docx.Document(self.fullpath)

    def _get_word_text(self):
        # load docx
        word_text = []
        for para in self.doc.paragraphs:
            word_text.append(para.text)

        # table parsing
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        word_text.append(para.text)

        # heading parsing
        for content in self.doc.paragraphs:
            if content.style.name == 'Heading 1' or content.style.name == 'Heading 2' or content.style.name == 'Heading 3':
                word_text.append(content.text)

        return word_text

    def get_text_docx_main_str(self):
        word_text = self._get_word_text()
        res = '\n'.join(word_text)
        return res

    def get_text_docx_main_list(self):
        return [i for i in self._get_word_text() if i != '' and i != '\n']

    def get_text_docx_main_with_filename(self):
        return [self.filename, '\n'.join(self._get_word_text())]


if __name__ == '__main__':

    from doc_miner.common_modules.common import *

    # resources 경로가져오기
    fm = 'docx'
    path = os.getcwd()
    rsc_path = OperatingSystem.get_parent_path(path)
    sample_path = os.path.join(rsc_path, 'resources/sample_doc')
    docx_files = OperatingSystem.sperate_file_extension(check_dir=sample_path)[fm]

    # .docx text 파싱
    sample_docx_filename = docx_files[0]
    ins_docx = GetDocxText(filepath=sample_path, filename=sample_docx_filename)

    res_list = ins_docx.get_text_docx_main_list()
    print(res_list)

    res_str = ins_docx.get_text_docx_main_str()
    print(res_str)

